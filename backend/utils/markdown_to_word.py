"""
Markdown 转 Word 工具
将翻译结果的 Markdown 格式转换为 Word 文档
"""

import re
import io
from typing import List, Tuple
from loguru import logger

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("[WARNING] python-docx 未安装，Word 导出功能将不可用")


class MarkdownToWordConverter:
    """
    Markdown 转 Word 转换器
    ��门用于处理公文翻译的对照格式
    """

    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安装，请运行: pip install python-docx")

    def convert_report(self, markdown_text: str) -> bytes:
        """
        转换纯中文报告（国别报告、季度报告）

        Args:
            markdown_text: Markdown 文本

        Returns:
            Word 文档字节流
        """
        doc = Document()

        # 设置文档属性
        self._setup_report_styles(doc)

        # 解析并处理内容（支持表格）
        lines = markdown_text.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i].rstrip()

            # 检测Markdown表格
            if self._is_table_line(line):
                # 收集所有表格行
                table_lines = []
                while i < len(lines) and self._is_table_line(lines[i].rstrip()):
                    table_lines.append(lines[i].rstrip())
                    i += 1

                # 添加表格到文档
                self._add_table_to_doc(doc, table_lines)
            else:
                # 普通行处理
                self._add_report_line(doc, line)
                i += 1

        # 保存到字节流
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()

    def _is_table_line(self, line: str) -> bool:
        """
        判断是否为Markdown表格行

        Args:
            line: 文本行

        Returns:
            是否为表格行
        """
        # 表格行包含 |
        return '|' in line and line.strip().startswith('|')

    def _add_table_to_doc(self, doc: Document, table_lines: List[str]):
        """
        添加Markdown表格到Word文档

        Args:
            doc: Word文档对象
            table_lines: 表格行列表
        """
        if not table_lines:
            return

        # 过滤掉分隔行（如 |---|---|）
        data_lines = [
            line for line in table_lines
            if not all(cell.strip().startswith('---') for cell in line.split('|')[1:-1])
        ]

        if not data_lines:
            return

        # 解析表格数据
        table_data = []
        for line in data_lines:
            # 移除首尾的 |
            cells = line.strip().strip('|').split('|')
            # 清理每个单元格
            cells = [cell.strip() for cell in cells]
            table_data.append(cells)

        if not table_data:
            return

        # 创建表格
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Table Grid'

        # 填充表格数据
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = cell_data

                    # 设置单元格字体
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '宋体'
                            run.font.size = Pt(10.5)
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                    # 表头样式（第一行）
                    if i == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(11)

    def _setup_report_styles(self, doc: Document):
        """
        设置报告文档样式

        Args:
            doc: Word 文档对象
        """
        # 设置默认字体（仿宋/宋体，公文标准）
        style = doc.styles['Normal']
        font = style.font
        font.name = '仿宋'
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)

        # 设置中文字体
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

        # 设置段落格式
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.5
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.first_line_indent = Inches(0.29)

    def _add_report_line(self, doc: Document, line: str):
        """
        添加报告行到文档

        Args:
            doc: Word 文档对象
            line: 文本行
        """
        line = line.rstrip()

        # 空行
        if not line:
            doc.add_paragraph()
            return

        # 一级标题 (#)
        if line.startswith('# '):
            title = line[2:].strip()
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            heading.runs[0].font.name = '宋体'
            heading.runs[0].font.size = Pt(16)
            heading.runs[0].font.bold = True
            heading.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            return

        # 二级标题 (##)
        if line.startswith('## '):
            title = line[3:].strip()
            heading = doc.add_heading(title, level=2)
            heading.runs[0].font.name = '黑体'
            heading.runs[0].font.size = Pt(14)
            heading.runs[0].font.bold = True
            heading.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            return

        # 三级标题 (###)
        if line.startswith('### '):
            title = line[4:].strip()
            heading = doc.add_heading(title, level=3)
            heading.runs[0].font.name = '楷体'
            heading.runs[0].font.size = Pt(13)
            heading.runs[0].font.bold = True
            heading.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            return

        # 加粗文本 **text**
        if line.startswith('**') and line.endswith('**'):
            text = line[2:-2]
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.name = '黑体'
            run.font.size = Pt(12)
            run.font.bold = True
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            return

        # 分隔线
        if line.startswith('---') or line.startswith('***'):
            para = doc.add_paragraph()
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = para.add_run('─' * 20)
            run.font.color.rgb = RGBColor(128, 128, 128)
            return

        # 数据来源标注 [Data Source: ...] 或 【数据来源：...】
        if '[Data Source:' in line or '【数据来源：' in line:
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            para.paragraph_format.first_line_indent = 0
            return

        # 普通段落
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.name = '仿宋'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    def convert(self, markdown_text: str) -> bytes:
        """
        将 Markdown 文本转换为 Word 文档字节流

        Args:
            markdown_text: Markdown 格式的文本

        Returns:
            Word 文档的字节流
        """
        doc = Document()

        # 设置文档默认字体
        self._setup_document_styles(doc)

        # 解析 Markdown 并添加到文档
        blocks = self._parse_markdown_blocks(markdown_text)

        for block_type, content in blocks:
            self._add_block_to_doc(doc, block_type, content)

        # 保存到字节流
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()

    def _setup_document_styles(self, doc: Document):
        """
        设置文档样式

        Args:
            doc: Word 文档对象
        """
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(10.5)
        font.color.rgb = RGBColor(0, 0, 0)

        # 设置中文字体
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 设置段落间距
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.5
        paragraph_format.space_before = Pt(6)
        paragraph_format.space_after = Pt(6)

    def _parse_markdown_blocks(self, markdown_text: str) -> List[Tuple[str, str]]:
        """
        解析 Markdown 文本为块列表

        Args:
            markdown_text: Markdown 文本

        Returns:
            (块类型, 内容) 的列表
            块类型: 'heading', 'chinese', 'english', 'separator'
        """
        blocks = []
        lines = markdown_text.split('\n')

        i = 0
        while i < len(lines):
            line = lines[i].strip()

            # 跳过空行
            if not line:
                i += 1
                continue

            # 标题 (# ## ###)
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                title_text = line.lstrip('#').strip()
                blocks.append(('heading', f"{level}|{title_text}"))
                i += 1
                continue

            # 分隔线 (---)
            if line.startswith('---') or line.startswith('***'):
                blocks.append(('separator', ''))
                i += 1
                continue

            # 中文译文标记
            if line.startswith('**中文译文**') or line.startswith('**译文**'):
                # 提取冒号后的内容
                if ':' in line or '：' in line:
                    content_start = max(line.find(':'), line.find('：')) + 1
                    content = line[content_start:].strip()
                    if content:
                        blocks.append(('chinese', content))
                    else:
                        # 内容在下一行
                        i += 1
                        if i < len(lines):
                            blocks.append(('chinese', lines[i].strip()))
                i += 1
                continue

            # 英文原文标记
            if line.startswith('**英文原文**') or line.startswith('**原文**'):
                # 提取冒号后的内容
                if ':' in line or '：' in line:
                    content_start = max(line.find(':'), line.find('：')) + 1
                    content = line[content_start:].strip()
                    if content:
                        blocks.append(('english', content))
                    else:
                        # 内容在下一行
                        i += 1
                        if i < len(lines):
                            blocks.append(('english', lines[i].strip()))
                i += 1
                continue

            # 普通文本段落（可能是中英文对照的延续）
            # 判断是否为中文或英文
            if self._is_chinese_text(line):
                blocks.append(('chinese', line))
            else:
                blocks.append(('english', line))

            i += 1

        return blocks

    def _is_chinese_text(self, text: str) -> bool:
        """
        判断文本是否主要为中文

        Args:
            text: 待判断的文本

        Returns:
            是否为中文
        """
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        total_chars = len(re.findall(r'[^\s]', text))
        if total_chars == 0:
            return False
        return chinese_chars / total_chars > 0.3

    def _add_block_to_doc(self, doc: Document, block_type: str, content: str):
        """
        将块添加到 Word 文档

        Args:
            doc: Word 文档对象
            block_type: 块类型
            content: 块内容
        """
        if block_type == 'separator':
            # 添加分隔线
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run('─' * 30)
            run.font.color.rgb = RGBColor(128, 128, 128)

        elif block_type == 'heading':
            # 添加标题
            level, title = content.split('|', 1)
            level = min(int(level), 3)  # 最多3级

            heading_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3'}
            doc.add_heading(title, level=level)

        elif block_type == 'chinese':
            # 添加中文译文
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"【译文】{content}")
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.bold = False

        elif block_type == 'english':
            # 添加英文原文
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"【原文】{content}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(80, 80, 80)
            run.font.italic = True

    def convert_bilingual(self, markdown_text: str) -> bytes:
        """
        转换双语文档（更精美的格式）

        Args:
            markdown_text: Markdown 文本

        Returns:
            Word 文档字节流
        """
        doc = Document()

        # 设置文档属性
        core_props = doc.core_properties
        core_props.title = "公文翻译"
        core_props.comments = "由 AI 公文翻译系统生成"

        # 设置默认字体
        self._setup_document_styles(doc)

        # 添加标题
        title = doc.add_heading('公文翻译结果', level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 解析并添加内容
        blocks = self._parse_markdown_blocks(markdown_text)

        # 逐块处理
        current_chinese = []
        current_english = []

        for block_type, content in blocks:
            if block_type == 'separator':
                # 遇到分隔符，输出当前累积的内容
                if current_chinese or current_english:
                    self._add_bilingual_block(doc, current_chinese, current_english)
                    current_chinese = []
                    current_english = []
                # 添加空行
                doc.add_paragraph()

            elif block_type == 'heading':
                # 先输出之前累积的内容
                if current_chinese or current_english:
                    self._add_bilingual_block(doc, current_chinese, current_english)
                    current_chinese = []
                    current_english = []
                # 添加标题
                level, title = content.split('|', 1)
                doc.add_heading(title, level=min(int(level), 2))

            elif block_type == 'chinese':
                current_chinese.append(content)

            elif block_type == 'english':
                current_english.append(content)

        # 输出最后的内容
        if current_chinese or current_english:
            self._add_bilingual_block(doc, current_chinese, current_english)

        # 保存到字节流
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()

    def _add_bilingual_block(self, doc: Document, chinese_lines: List[str], english_lines: List[str]):
        """
        添加双语文本块

        Args:
            doc: Word 文档对象
            chinese_lines: 中文行列表
            english_lines: 英文行列表
        """
        # 中文段落
        if chinese_lines:
            para = doc.add_paragraph()
            run = para.add_run(''.join(chinese_lines))
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 英文段落
        if english_lines:
            para = doc.add_paragraph()
            run = para.add_run(''.join(english_lines))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.font.italic = True


# 单例实例
markdown_to_word_converter = MarkdownToWordConverter() if DOCX_AVAILABLE else None


def convert_markdown_to_word(markdown_text: str, use_bilingual_format: bool = True, document_type: str = "translation") -> bytes:
    """
    将 Markdown 文本转换为 Word 文档

    Args:
        markdown_text: Markdown 格式的文本
        use_bilingual_format: 是否使用双语对照格式
        document_type: 文档类型 ("translation"=翻译, "report"=报告)

    Returns:
        Word 文档的字节流
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx 未安装，请运行: pip install python-docx")

    converter = MarkdownToWordConverter()

    if document_type == "report":
        return converter.convert_report(markdown_text)
    elif use_bilingual_format:
        return converter.convert_bilingual(markdown_text)
    else:
        return converter.convert(markdown_text)
