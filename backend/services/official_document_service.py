"""
公文库服务
"""

import os
import hashlib
import uuid
from typing import List, Tuple, Optional
from datetime import datetime
from pathlib import Path
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, and_, or_, desc, asc, func
from fastapi import HTTPException, status

from models.official_document import OfficialDocument
from models.user import User
from schemas.official_document import (
    OfficialDocumentCreate,
    OfficialDocumentUpdate,
    OfficialDocumentListQuery,
)
from core.config import settings


class OfficialDocumentService:
    """公文库服务"""

    def __init__(self):
        self.storage_dir = Path(settings.STORAGE_ROOT) / "official_documents"
        self.storage_dir.mkdir(parents=True, exist_ok=True)

    async def list_documents(
        self,
        query_params: OfficialDocumentListQuery,
        owner_id: str,
        db: AsyncSession,
    ) -> Tuple[List[OfficialDocument], int]:
        """
        获取公文列表

        Args:
            query_params: 查询参数
            owner_id: 用户 ID
            db: 数据库会话

        Returns:
            (公文列表, 总数)
        """
        # 构建查询
        query = select(OfficialDocument).where(OfficialDocument.owner_id == owner_id)

        # 筛选条件
        filters = []

        if query_params.source:
            filters.append(OfficialDocument.source == query_params.source)

        if query_params.keyword:
            keyword_filter = or_(
                OfficialDocument.title.ilike(f"%{query_params.keyword}%"),
                OfficialDocument.description.ilike(f"%{query_params.keyword}%"),
            )
            filters.append(keyword_filter)

        if query_params.is_verified is not None:
            filters.append(OfficialDocument.is_verified == query_params.is_verified)

        if filters:
            query = query.where(and_(*filters))

        # 排序
        sort_column = getattr(OfficialDocument, query_params.sort_by, OfficialDocument.created_at)
        if query_params.sort_order == "asc":
            query = query.order_by(asc(sort_column))
        else:
            query = query.order_by(desc(sort_column))

        # 获取总数
        count_query = select(func.count()).select_from(query.subquery())
        total_result = await db.execute(count_query)
        total = total_result.scalar()

        # 分页
        offset = (query_params.page - 1) * query_params.page_size
        query = query.offset(offset).limit(query_params.page_size)

        # 执行查询
        result = await db.execute(query)
        documents = result.scalars().all()

        return documents, total

    async def get_document(
        self,
        document_id: str,
        owner_id: str,
        db: AsyncSession,
    ) -> Optional[OfficialDocument]:
        """
        获取单个公文

        Args:
            document_id: 公文 ID
            owner_id: 用户 ID
            db: 数据库会话

        Returns:
            公文对象或 None
        """
        query = select(OfficialDocument).where(
            and_(
                OfficialDocument.id == document_id,
                OfficialDocument.owner_id == owner_id,
            )
        )

        result = await db.execute(query)
        return result.scalar_one_or_none()

    async def create_document(
        self,
        document_data: OfficialDocumentCreate,
        file_content: bytes,
        filename: str,
        mime_type: str,
        owner_id: str,
        db: AsyncSession,
    ) -> OfficialDocument:
        """
        创建公文

        Args:
            document_data: 公文数据
            file_content: 文件内容
            filename: 文件名
            mime_type: MIME 类型
            owner_id: 用户 ID
            db: 数据库会话

        Returns:
            创建的公文对象
        """
        # 1. 计算文件哈希
        file_hash = hashlib.sha256(file_content).hexdigest()

        # 2. 检查是否已存在
        existing_query = select(OfficialDocument).where(
            and_(
                OfficialDocument.file_hash == file_hash,
                OfficialDocument.owner_id == owner_id,
            )
        )
        existing_result = await db.execute(existing_query)
        existing_doc = existing_result.scalar_one_or_none()

        if existing_doc:
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail="该文件已存在",
            )

        # 3. 确定文档类型
        document_type = self._get_document_type(filename, mime_type)

        # 4. 保存文件
        file_extension = Path(filename).suffix
        storage_filename = f"{uuid.uuid4()}{file_extension}"
        file_path = self.storage_dir / storage_filename

        with open(file_path, "wb") as f:
            f.write(file_content)

        # 5. 提取文本内容
        content, content_preview, page_count = await self._extract_content(
            file_path=file_path,
            document_type=document_type,
        )

        # 6. 创建公文记录
        import time
        current_time = int(time.time())

        document = OfficialDocument(
            id=str(uuid.uuid4()),
            title=document_data.title,
            original_filename=filename,
            document_type=document_type,
            description=document_data.description,
            source=document_data.source,
            file_path=str(file_path.relative_to(settings.STORAGE_ROOT)),
            file_size=len(file_content),
            mime_type=mime_type,
            file_hash=file_hash,
            content=content,
            content_preview=content_preview,
            page_count=page_count,
            is_verified=document_data.is_verified,
            owner_id=owner_id,
            created_at=current_time,
            updated_at=current_time,
        )

        db.add(document)
        await db.commit()
        await db.refresh(document)

        return document

    async def update_document(
        self,
        document_id: str,
        document_data: OfficialDocumentUpdate,
        owner_id: str,
        db: AsyncSession,
    ) -> Optional[OfficialDocument]:
        """
        更新公文

        Args:
            document_id: 公文 ID
            document_data: 更新数据
            owner_id: 用户 ID
            db: 数据库会话

        Returns:
            更新后的公文对象或 None
        """
        # 获取公文
        document = await self.get_document(document_id, owner_id, db)

        if not document:
            return None

        # 更新字段
        update_data = document_data.model_dump(exclude_unset=True)

        for field, value in update_data.items():
            setattr(document, field, value)

        # 如果设置为已审核，记录审核信息
        if document_data.is_verified and not document.is_verified:
            document.verified_at = datetime.now()

        await db.commit()
        await db.refresh(document)

        return document

    async def delete_document(
        self,
        document_id: str,
        owner_id: str,
        db: AsyncSession,
    ) -> bool:
        """
        删除公文

        Args:
            document_id: 公文 ID
            owner_id: 用户 ID
            db: 数据库会话

        Returns:
            是否删除成功
        """
        # 获取公文
        document = await self.get_document(document_id, owner_id, db)

        if not document:
            return False

        # 删除文件
        file_path = Path(settings.STORAGE_ROOT) / document.file_path
        if file_path.exists():
            file_path.unlink()

        # 删除缩略图
        if document.thumbnail_path:
            thumbnail_path = Path(settings.STORAGE_ROOT) / document.thumbnail_path
            if thumbnail_path.exists():
                thumbnail_path.unlink()

        # 删除数据库记录
        await db.delete(document)
        await db.commit()

        return True

    def _get_document_type(self, filename: str, mime_type: str) -> str:
        """
        根据文件名和 MIME 类型确定文档类型

        Args:
            filename: 文件名
            mime_type: MIME 类型

        Returns:
            文档类型字符串
        """
        # 根据 MIME 类型判断
        if mime_type == "application/pdf":
            return "pdf"
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return "docx"
        elif mime_type == "text/plain":
            return "txt"
        elif mime_type in ["text/markdown", "text/x-markdown"]:
            return "md"

        # 根据文件扩展名判断
        ext = Path(filename).suffix.lower()
        if ext == ".pdf":
            return "pdf"
        elif ext == ".docx":
            return "docx"
        elif ext in [".txt", ".text"]:
            return "txt"
        elif ext in [".md", ".markdown"]:
            return "md"

        return "unknown"

    async def _extract_content(
        self,
        file_path: Path,
        document_type: str,
    ) -> Tuple[str, str, Optional[int]]:
        """
        提取文档内容

        Args:
            file_path: 文件路径
            document_type: 文档类型

        Returns:
            (完整内容, 内容预览, 页数)
        """
        try:
            if document_type == "pdf":
                return await self._extract_pdf_content(file_path)
            elif document_type == "docx":
                return await self._extract_docx_content(file_path)
            elif document_type in ["txt", "md"]:
                return await self._extract_text_content(file_path)
            else:
                return "", None, None
        except Exception as e:
            print(f"Error extracting content: {e}")
            return "", None, None

    async def _extract_pdf_content(self, file_path: Path) -> Tuple[str, str, Optional[int]]:
        """提取 PDF 内容"""
        import fitz

        doc = fitz.open(str(file_path))
        text_content = []

        for page in doc:
            text_content.append(page.get_text())

        full_text = "\n\n".join(text_content)
        preview = full_text[:500] if full_text else None

        doc.close()

        return full_text, preview, len(doc) if hasattr(doc, "__len__") else None

    async def _extract_docx_content(self, file_path: Path) -> Tuple[str, str, None]:
        """提取 DOCX 内容"""
        from docx import Document

        doc = Document(str(file_path))

        # 提取所有段落
        all_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text)

        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                if row_text.strip():
                    all_text.append(row_text)

        full_text = "\n\n".join(all_text)
        preview = full_text[:500] if full_text else None

        print(f"提取完成 - 段落数: {len(doc.paragraphs)}, 表格数: {len(doc.tables)}, 总字符数: {len(full_text)}")

        return full_text, preview, None

    async def _extract_text_content(self, file_path: Path) -> Tuple[str, str, None]:
        """提取文本内容"""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        preview = content[:500] if content else None

        return content, preview, None


# 创建全局服务实例
official_document_service = OfficialDocumentService()
